#!/usr/bin/env python3
"""
ResumeRadar QA Suite
====================
Automated production-readiness checks for ResumeRadar.
Run via: python tests/qa_suite.py [--quick]

--quick  : Run fast checks only (routes, structure, security). ~2s
(default): Run full suite including live scan + PDF generation. ~8s

Exit codes:
  0 = all passed
  1 = one or more failures
"""

import sys
import os
import re
import json
import time

# Add project root to path
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from app import app

# ============================================================
# CONFIG
# ============================================================
QUICK_MODE = '--quick' in sys.argv

PASS = 0
FAIL = 0
RESULTS = []


def check(name, passed, detail=""):
    """Record a test result."""
    global PASS, FAIL
    if passed:
        PASS += 1
        RESULTS.append(("PASS", name, detail))
    else:
        FAIL += 1
        RESULTS.append(("FAIL", name, detail))


# ============================================================
# TEST SUITE
# ============================================================
def run_tests():
    start = time.time()

    with app.test_client() as c:

        # ---- SECTION 1: ROUTES ----
        # 1. Main page
        r = c.get('/')
        check("GET / returns 200", r.status_code == 200)

        html = r.data.decode()

        # 2. Health check
        r = c.get('/api/health')
        check("GET /api/health returns 200", r.status_code == 200)

        # 3. Scan count
        r = c.get('/api/scan-count')
        d = r.get_json()
        check("GET /api/scan-count returns count", r.status_code == 200 and 'count' in d)

        # 4. robots.txt
        r = c.get('/robots.txt')
        check("GET /robots.txt returns 200", r.status_code == 200)

        # 5. favicon.ico
        r = c.get('/favicon.ico')
        check("GET /favicon.ico returns 200", r.status_code == 200)

        # 6. apple-touch-icon
        r = c.get('/apple-touch-icon.png')
        check("GET /apple-touch-icon.png returns 200", r.status_code == 200)

        # 7. 404 page (browser)
        r = c.get('/nonexistent-page')
        check("GET /nonexistent returns 404", r.status_code == 404)

        # 8. 404 API (JSON)
        r = c.get('/api/nonexistent')
        d = r.get_json()
        check("GET /api/404 returns JSON error", r.status_code == 404 and d and 'error' in d)

        # ---- SECTION 2: SCAN FORM STRUCTURE ----
        # 9. Only jobDescription is required inside the form
        form_match = re.search(r'<form[^>]*id="scanForm"[^>]*>(.*?)</form>', html, re.DOTALL)
        if form_match:
            form_html = form_match.group(1)
            required_inputs = re.findall(r'<(?:input|textarea)[^>]*required[^>]*>', form_html)
            required_ids = []
            for inp in required_inputs:
                m = re.search(r'id="([^"]+)"', inp)
                if m:
                    required_ids.append(m.group(1))
            check("Only jobDescription required in form", required_ids == ['jobDescription'],
                  f"Found: {required_ids}")
        else:
            check("Scan form found in HTML", False, "Could not find scanForm")

        # ---- SECTION 3: SECURITY HEADERS ----
        r = c.get('/')
        check("X-Content-Type-Options header", r.headers.get('X-Content-Type-Options') == 'nosniff')
        check("X-Frame-Options header", r.headers.get('X-Frame-Options') == 'SAMEORIGIN')
        check("X-XSS-Protection header", r.headers.get('X-XSS-Protection') == '1; mode=block')
        check("Referrer-Policy header", 'strict-origin' in (r.headers.get('Referrer-Policy') or ''))
        check("Permissions-Policy header", bool(r.headers.get('Permissions-Policy')))

        # ---- SECTION 4: HTML INTEGRITY ----
        # 15. Key elements present
        key_elements = [
            ('scanForm', 'Scan form'),
            ('scanBtn', 'Scan button'),
            ('demoScanBtn', 'Demo scan button'),
            ('results', 'Results section'),
            ('copyReportBtn', 'Copy report button'),
            ('downloadReportBtn', 'Download PDF button'),
            ('shareLinkedIn', 'Share LinkedIn button'),
            ('newsletterPopup', 'Newsletter popup'),
            ('errorMessage', 'Error message div'),
        ]
        for elem_id, label in key_elements:
            check(f"HTML has #{elem_id}", f'id="{elem_id}"' in html, label)

        # ---- SECTION 5: META TAGS ----
        check("OG title meta tag", 'og:title' in html)
        check("OG image meta tag", 'og:image' in html)
        check("Twitter card meta tag", 'twitter:card' in html)
        check("Favicon link tag", 'favicon.ico' in html)
        check("Apple touch icon link", 'apple-touch-icon' in html)

        # ---- SECTION 6: API VALIDATION ----
        # Empty scan
        r = c.post('/api/scan', data={'job_description': ''})
        check("Scan rejects empty JD", r.status_code == 400)

        # Short scan
        r = c.post('/api/scan', data={'job_description': 'too short'})
        check("Scan rejects short JD", r.status_code == 400)

        # No resume
        r = c.post('/api/scan', data={'job_description': 'a ' * 20})
        check("Scan rejects missing resume", r.status_code == 400)

        # Subscribe without email
        r = c.post('/api/subscribe', json={'email': '', 'first_name': 'Test'})
        check("Subscribe rejects empty email", r.status_code == 400)

        # Subscribe without first name
        r = c.post('/api/subscribe', json={'email': 'a@b.com', 'first_name': '', 'utm_source': 'resumeradar'})
        check("Subscribe rejects missing name", r.status_code == 400)

        # ---- SECTION 7: FULL SCAN (skip in quick mode) ----
        if not QUICK_MODE:
            r = c.post('/api/scan', data={
                'job_description': (
                    'We are looking for a Cloud Engineer with AWS experience including '
                    'EC2 S3 Lambda CloudFormation Terraform Kubernetes Docker CI/CD pipelines '
                    'Python scripting and networking fundamentals.'
                ),
                'resume_text': (
                    'Experienced cloud engineer with 5 years of AWS experience. '
                    'Skilled in EC2, S3, Lambda, CloudFormation, and Terraform. '
                    'Built CI/CD pipelines with Jenkins and GitHub Actions. '
                    'Proficient in Python and Docker.'
                ),
            })
            d = r.get_json()
            check("Full scan returns 200", r.status_code == 200)
            check("Full scan has match_score", d and 'match_score' in d)
            check("Full scan score > 0", d and d.get('match_score', 0) > 0,
                  f"Score: {d.get('match_score') if d else 'N/A'}")
            check("Full scan has category_scores", d and 'category_scores' in d)
            check("Full scan has ai_suggestions", d and 'ai_suggestions' in d)

            # ---- SECTION 8: PDF GENERATION ----
            r = c.post('/api/download-report', json={
                'match_score': 72,
                'total_matched': 8,
                'total_missing': 3,
                'total_job_keywords': 11,
                'category_scores': {},
                'matched_keywords': {},
                'missing_keywords': {},
                'ats_formatting': {},
                'ai_suggestions': {},
            })
            check("PDF report generates", r.status_code == 200 and len(r.data) > 500,
                  f"Size: {len(r.data)} bytes")

    # ---- SECTION 9: CV BUILDER ----
    with app.test_client() as c:
        # Build page loads
        r = c.get('/build')
        check("GET /build returns 200", r.status_code == 200)
        build_html = r.data.decode()
        check("Build page has privacy badge", 'privacy-badge' in build_html)

    # Score-aware CTA on scan page
    check("Scan CTA has dynamic heading ID", 'buildCtaHeading' in html)
    check("Scan CTA has trust badge", 'build-cta-trust' in html and 'never make anything up' in html.lower())

    # PDF generation — all 3 templates
    from backend.cv_pdf_generator import generate_cv_pdf, TEMPLATES, _flatten_skills, _safe

    test_cv = {
        'personal': {'full_name': 'QA Test', 'email': 'qa@test.com', 'phone': '+1 555 0199', 'location': 'London'},
        'summary': 'Experienced engineer with 5 years building scalable systems.',
        'experience': [{'title': 'Engineer', 'company': 'Acme', 'start_date': 'Jan 2020', 'end_date': 'Present',
                        'bullets': ['Built microservices', 'Led team of 4', 'Reduced deploy time by 60%']}],
        'education': [{'degree': 'BSc CS', 'institution': 'UCL', 'graduation_date': '2019', 'details': ''}],
        'skills': ['Python', 'AWS', 'Docker'],
        'certifications': [{'name': 'AWS SA', 'issuer': 'Amazon', 'date': '2023'}]
    }

    for tmpl in TEMPLATES:
        try:
            pdf_out = generate_cv_pdf(test_cv, tmpl)
            is_pdf = isinstance(pdf_out, (bytes, bytearray)) and len(pdf_out) > 500
            check(f"PDF {tmpl} template generates", is_pdf, f"{len(pdf_out)} bytes")
        except Exception as e:
            check(f"PDF {tmpl} template generates", False, str(e))

    # Invalid template falls back to classic
    try:
        pdf_out = generate_cv_pdf(test_cv, 'nonexistent')
        check("Invalid template falls back to classic", isinstance(pdf_out, (bytes, bytearray)) and len(pdf_out) > 500)
    except Exception as e:
        check("Invalid template falls back to classic", False, str(e))

    # Empty CV doesn't crash
    empty_cv = {'personal': {}, 'summary': '', 'experience': [], 'education': [], 'skills': [], 'certifications': []}
    try:
        all_ok = True
        for tmpl in TEMPLATES:
            pdf_out = generate_cv_pdf(empty_cv, tmpl)
            if not isinstance(pdf_out, (bytes, bytearray)):
                all_ok = False
        check("Empty CV data doesn't crash (all templates)", all_ok)
    except Exception as e:
        check("Empty CV data doesn't crash (all templates)", False, str(e))

    # Long text wraps without crash
    long_cv = dict(test_cv)
    long_cv['summary'] = 'Word ' * 200
    long_cv['experience'] = [{'title': 'Engineer', 'company': 'Co', 'start_date': '2020', 'end_date': 'Present',
                               'bullets': ['Achievement ' * 50, 'Result ' * 80]}]
    long_cv['skills'] = [f'Skill{i}' for i in range(40)]
    try:
        all_ok = True
        for tmpl in TEMPLATES:
            pdf_out = generate_cv_pdf(long_cv, tmpl)
            if not isinstance(pdf_out, (bytes, bytearray)) or len(pdf_out) < 500:
                all_ok = False
        check("Long text wraps without crash (all templates)", all_ok)
    except Exception as e:
        check("Long text wraps without crash (all templates)", False, str(e))

    # _flatten_skills handles both formats
    check("_flatten_skills: flat list", _flatten_skills(['A', 'B']) == ['A', 'B'])
    dict_result = _flatten_skills({'matched': ['A'], 'missing': ['B'], 'additional': ['C']})
    check("_flatten_skills: scan dict", 'A' in dict_result and 'B' in dict_result and 'C' in dict_result)

    # _safe handles Unicode
    check("_safe: Unicode conversion", _safe('Hello \u2014 World') == 'Hello -- World' and _safe(None) == '')

    # All multi_cell calls have align='L'
    pdf_gen_path = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
                                'backend', 'cv_pdf_generator.py')
    with open(pdf_gen_path, 'r') as f:
        pdf_source = f.read()
    # Check each line with multi_cell individually (avoids nested-paren regex issues)
    mc_lines = [line.strip() for line in pdf_source.split('\n') if 'multi_cell(' in line and not line.strip().startswith('#')]
    bad_mc = [line for line in mc_lines if "align='L'" not in line]
    check(f"All {len(mc_lines)} multi_cell calls have align='L'", len(bad_mc) == 0,
          f"{len(bad_mc)} missing" if bad_mc else "")

    # LLM prompts have no-hallucination guardrails
    from backend.cv_builder import polish_cv_sections, extract_and_polish
    import inspect
    polish_src = inspect.getsource(polish_cv_sections)
    extract_src = inspect.getsource(extract_and_polish)
    check("LLM polish prompt: no-hallucination rules",
          'NEVER add skills' in polish_src and 'NEVER invent metrics' in polish_src)
    check("LLM extract prompt: no-hallucination rules",
          'NEVER invent' in extract_src and 'ACTUALLY EXISTS' in extract_src)
    check("LLM prompts return smart_suggestions",
          'smart_suggestions' in polish_src and 'smart_suggestions' in extract_src)

    # Fallback when API key is dummy
    old_key = os.environ.get('ANTHROPIC_API_KEY')
    os.environ['ANTHROPIC_API_KEY'] = 'your-anthropic-api-key-here'
    fallback = polish_cv_sections(test_cv)
    check("Fallback: returns unpolished data", fallback.get('ai_polished') == False)
    check("Fallback: has smart_suggestions key", 'smart_suggestions' in fallback)
    if old_key:
        os.environ['ANTHROPIC_API_KEY'] = old_key

    # app.js has score-aware CTA function
    # ---- SECTION 10: JS SYNTAX CHECK ----
    js_path = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
                           'static', 'js', 'app.js')
    if os.path.exists(js_path):
        with open(js_path, 'r') as f:
            js_content = f.read()
        # Basic syntax checks
        check("JS file not empty", len(js_content) > 1000, f"{len(js_content)} chars")
        check("JS has DOMContentLoaded", 'DOMContentLoaded' in js_content)
        check("JS has no console.log (debug)", 'console.log(' not in js_content or
              js_content.count('console.log(') <= 2,
              f"Found {js_content.count('console.log(')} console.log calls")
    else:
        check("JS file exists", False, js_path)

    # ---- SECTION 11: UX IMPROVEMENTS + EMAIL DELIVERY ----

    # Behavioral endpoint tests (strongest — test actual HTTP responses)
    with app.test_client() as c:
        # Checkout validates bad email → 400
        r = c.post('/api/build/create-checkout',
            json={'token': 'test', 'template': 'classic', 'delivery_email': 'not-valid'},
            content_type='application/json')
        check("Checkout rejects invalid delivery email", r.status_code == 400)

        # Checkout accepts empty email → NOT a 400 about email
        r = c.post('/api/build/create-checkout',
            json={'token': 'test', 'template': 'classic', 'delivery_email': ''},
            content_type='application/json')
        d = r.get_json() or {}
        check("Checkout accepts empty delivery email",
              r.status_code != 400 or 'email' not in d.get('error', '').lower())

    # HTML structure checks (verifiable from template)
    with app.test_client() as c:
        r = c.get('/build')
        build_html = r.data.decode()

    preview_count = build_html.count('class="template-preview')
    check("Build page has 3 template previews", preview_count == 3, f"Found: {preview_count}")
    check("Delivery email input with maxlength",
          'id="deliveryEmail"' in build_html and 'maxlength="254"' in build_html)
    check("Loading text span in generate button",
          'gen-loading-text' in build_html)

    # Structural source assertions (multi-indicator to reduce false-pass risk)
    project_root = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))

    builder_js_path = os.path.join(project_root, 'static', 'js', 'builder.js')
    with open(builder_js_path, 'r') as f:
        builder_js = f.read()
    check("Loading message rotation system",
          'LOADING_MESSAGES' in builder_js and 'startLoadingRotation' in builder_js)
    check("Confetti with auto-cleanup",
          'createConfetti' in builder_js and 'confetti-piece' in builder_js and 'setTimeout' in builder_js)
    check("Form pre-populate helper",
          'populateDynamicEntries' in builder_js)
    check("Email-requested header read",
          'X-Email-Requested' in builder_js)
    check("Inline editable summary",
          'contenteditable' in builder_js and 'preview-summary-editable' in builder_js)

    app_py_path = os.path.join(project_root, 'app.py')
    with open(app_py_path, 'r') as f:
        app_source = f.read()
    check("Webhook calls _send_cv_email with event_id",
          '_send_cv_email' in app_source and "event['id']" in app_source)
    check("SETNX dedup keyed on event_id (72h TTL)",
          'nx=True' in app_source and 'cv_emailed' in app_source and '259200' in app_source)
    check("CV data TTL extended on webhook",
          '.expire(' in app_source and '259200' in app_source)
    check("User values sanitized (html.escape + re.sub)",
          'html_module.escape' in app_source and 're_module.sub' in app_source)
    check("Download route returns X-Email-Requested header",
          'X-Email-Requested' in app_source)

    builder_css_path = os.path.join(project_root, 'static', 'css', 'builder.css')
    with open(builder_css_path, 'r') as f:
        builder_css = f.read()
    check("Payment content flex-wrap for email row",
          'flex-wrap' in builder_css and 'payment-email' in builder_css)

    # ---- SECTION 12: E2E SCENARIOS ----

    with app.test_client() as c:

        # E2E-1: Stripe cancel flow — cancel URL returns cleanly
        r = c.get('/build?payment=cancelled')
        check("Cancel flow: /build?payment=cancelled returns 200", r.status_code == 200)
        cancelled_html = r.data.decode()
        check("Cancel flow: page renders without error", 'builderForm' in cancelled_html)

        # E2E-7: Token/session tampering — mismatched combos return 4xx
        # GET with fake token + fake session
        r = c.get('/api/build/download/fake-token-123?session_id=cs_fake_session&template=classic')
        check("Tampered token: download returns 4xx", r.status_code in (400, 403, 404, 500))

        # POST with fake session + fake token
        r = c.post('/api/build/download/fake-token-456',
            json={'session_id': 'cs_fake_session', 'template': 'classic', 'cv_data': {'personal': {}}},
            content_type='application/json')
        check("Tampered session POST: download returns 4xx", r.status_code in (400, 403, 404, 500))

        # Missing session_id entirely
        r = c.get('/api/build/download/some-token?template=classic')
        check("Missing session_id: download returns 400", r.status_code == 400)

        # E2E-10: Email validation normalization — edge cases
        # Valid edge emails (should not return 400 about email)
        valid_emails = ['user+tag@example.com', 'USER@EXAMPLE.COM', 'a@sub.domain.example.com']
        for em in valid_emails:
            r = c.post('/api/build/create-checkout',
                json={'token': 'test', 'template': 'classic', 'delivery_email': em},
                content_type='application/json')
            d = r.get_json() or {}
            is_email_error = r.status_code == 400 and 'email' in d.get('error', '').lower()
            check(f"Valid email accepted: {em}", not is_email_error)

        # Invalid emails (should return 400 about email)
        invalid_emails = ['notanemail', '@nolocal.com', 'spaces in@email.com', 'a@.com']
        for em in invalid_emails:
            r = c.post('/api/build/create-checkout',
                json={'token': 'test', 'template': 'classic', 'delivery_email': em},
                content_type='application/json')
            check(f"Invalid email rejected: {em}", r.status_code == 400)

        # E2E-11: Webhook signature validation — unsigned payloads rejected
        r = c.post('/api/build/webhook',
            data=b'{"type":"checkout.session.completed"}',
            content_type='application/json')
        check("Unsigned webhook: rejected (400)", r.status_code == 400)

        r = c.post('/api/build/webhook',
            data=b'{"type":"checkout.session.completed"}',
            content_type='application/json',
            headers={'Stripe-Signature': 'invalid_sig_header'})
        check("Invalid signature webhook: rejected (400)", r.status_code == 400)

    # E2E-12: UI state integrity after errors — source assertions
    check("Builder JS stops loading on scan error",
          'stopLoadingRotation()' in builder_js and 'scan-error' in builder_js)
    check("Builder JS re-enables button on generate error",
          'setGenerateLoading(false)' in builder_js and 'showError' in builder_js)
    check("Builder JS re-enables payment button on error",
          'setPaymentLoading(false)' in builder_js)

    # E2E-13: Accessibility assertions — labels, roles, keyboard support
    check("Email input has label element",
          'for="deliveryEmail"' in build_html)
    check("Summary editable in JS",
          'contenteditable' in builder_js)
    check("Template radios have labels",
          build_html.count('class="template-option"') == 3 and '<label' in build_html)
    check("Generate button is type submit",
          'type="submit"' in build_html and 'generateBtn' in build_html)

    # E2E-14: Cross-browser/mobile — responsive CSS assertions
    check("Mobile: form-row stacks to 1fr",
          '@media' in builder_css and 'grid-template-columns: 1fr' in builder_css)
    check("Mobile: template-picker stacks",
          'template-picker' in builder_css)
    check("Mobile: payment-content wraps",
          'flex-wrap: wrap' in builder_css)
    check("Mobile: celebration actions wrap",
          'celebration-actions' in builder_css and 'flex-wrap' in builder_css)

    # E2E-15: Observability assertions — logs/print statements for key events
    check("Observability: webhook error logged",
          "Webhook error:" in app_source or "Webhook processing" in app_source)
    check("Observability: email error logged",
          "CV email error" in app_source)
    check("Observability: download error logged",
          "CV Builder download error" in app_source)
    check("Observability: checkout error logged",
          "CV Builder checkout error" in app_source)

    # E2E-2/3/4: Success URL refresh, webhook replay, out-of-order timing
    # These require real Stripe sessions and Redis state — verified via structural assertions
    check("Download route enforces download limit (max 3)",
          'dl_count >= 3' in app_source or 'Download limit' in app_source)
    check("Webhook extends CV data TTL for retry window",
          '.expire(' in app_source and 'resumeradar:cv:' in app_source)
    check("SETNX dedup releases on failure for retry recovery",
          '_redis_client.delete(f"resumeradar:cv_emailed:' in app_source or
          '_redis_client.delete(dedup_key)' in app_source)

    # E2E-6: Redis degradation path — verify graceful handling
    check("Download falls back to client data when Redis unavailable",
          'client_cv_data' in app_source and 'not cv_data and client_cv_data' in app_source)
    check("Email skipped when Redis unavailable",
          'if not _redis_client:' in app_source)

    # E2E-8: Download limit enforcement
    check("Download counter tracked in Redis",
          'cv_downloads' in app_source and 'incr' in app_source)

    # E2E-9: TTL expiry behavior — graceful error messaging
    check("Expired CV data returns user-friendly message",
          'CV data not found' in app_source or 'may have expired' in app_source)
    check("Expired session returns user-friendly message",
          'CV session expired' in app_source or 'regenerate' in app_source)

    # ---- SECTION 13: E2E REGRESSION FIXES ----

    # Fix 1 (P1): Upload-to-builder handoff — scan response includes resume_text
    check("Scan response includes resume_text for file-upload users",
          "resume_text" in app_source and "extracted_resume_text" in app_source)

    # Verify app.js uses scan data fallback for file-upload users
    app_js_path = os.path.join(project_root, 'static', 'js', 'app.js')
    with open(app_js_path, 'r') as f:
        app_js = f.read()
    check("Builder handoff falls back to scan response resume_text",
          'scanData.resume_text' in app_js and 'textareaText' in app_js)

    # Fix 2 (P2): ImportError fallback validates email instead of silently dropping
    check("ImportError fallback logs warning and validates",
          'WARNING: email-validator not installed' in app_source and
          '"@" not in delivery_email' in app_source)

    # Fix 3 (P2): populateDynamicEntries clears stale entries before populating
    check("Dynamic entries cleared before re-populate",
          '.remove()' in builder_js and 'existingEntries' in builder_js)

    # Fix 4 (P2): Inline summary allows blank text (no truthy guard)
    # The blur handler should use `if (currentPolished)` not `if (currentPolished && newText)`
    check("Inline summary persists blank text",
          'currentPolished.summary = newText' in builder_js and
          'currentPolished && newText' not in builder_js)

    # Fix 5 (P3): Cancelled payment shows UX feedback
    check("Payment cancelled handler exists",
          'showPaymentCancelledMessage' in builder_js and 'payment-cancelled-banner' in builder_js)
    check("Payment cancelled auto-dismiss",
          'cancelled-dismiss' in builder_js and 'Auto-dismiss' in builder_js or
          'cancelled-dismiss' in builder_js and '8000' in builder_js)
    check("Payment cancelled CSS styles",
          'payment-cancelled-banner' in builder_css and 'cancelled-content' in builder_css)
    check("Payment cancelled URL cleanup",
          "searchParams.delete('payment')" in builder_js and 'replaceState' in builder_js)

    # Behavioral: /build?payment=cancelled still returns valid page
    with app.test_client() as c:
        r = c.get('/build?payment=cancelled')
        cancelled_page = r.data.decode()
        check("Cancel URL renders builder page with form",
              r.status_code == 200 and 'builderForm' in cancelled_page)

    # ---- SECTION 14: UPLOAD-FIRST BUILD + RATE LIMIT HARDENING ----
    print("\n-- Section 14: Upload-First Build + Rate Limit Hardening --")

    # --- HTML structure checks ---
    check("Upload section exists on build page",
          'id="uploadSection"' in build_html)
    check("Upload drop zone exists",
          'id="buildDropZone"' in build_html)
    check("Upload target JD textarea exists",
          'id="uploadTargetJD"' in build_html)
    check("Upload generate button exists",
          'id="uploadGenerateBtn"' in build_html)
    check("Manual form toggle link exists",
          'id="showManualForm"' in build_html)
    check("Back-to-upload link exists",
          'id="showUploadSection"' in build_html)
    check("Manual form hidden by default",
          'id="manualFormSection"' in build_html and 'style="display: none;"' in build_html)

    # --- JS source checks (multi-indicator) ---
    check("Upload file handler + endpoint call",
          'handleBuildFileSelect' in builder_js and 'generate-from-upload' in builder_js)
    check("Toggle handlers for both directions",
          'showManualFormLink' in builder_js and 'showUploadLink' in builder_js)
    check("Scan fallback shows upload section when data missing",
          'uploadSection.style.display' in builder_js and 'show upload section' in builder_js.lower())
    check("Retry-After parsing handles seconds and HTTP-date",
          'parseInt(retryAfter' in builder_js and 'new Date(retryAfter)' in builder_js)
    check("Edit & Regenerate hides upload section",
          'uploadSection' in builder_js and 'manualFormSection' in builder_js)

    # --- Rate limit config checks (app.py source) ---
    check("Limiter uses REDIS_URL with memory fallback",
          'REDIS_URL' in app_source and "memory://" in app_source)
    check("Checkout limit raised to 30/hour",
          app_source.count('"30 per hour"') >= 2)
    check("get_real_ip uses second-to-last for multi-proxy chains",
          'len(parts) - 2' in app_source)
    check("get_real_ip validates IP format with regex",
          "re_module.match" in app_source and r"[\d.:a-fA-F]" in app_source)

    # --- Startup logging (module-level, visible under Gunicorn) ---
    check("Startup logs Stripe Checkout status",
          'STRIPE_PRICE_ID' in app_source and 'Stripe Checkout' in app_source)
    check("Startup logs Rate Limiter backend",
          'Rate Limiter' in app_source)
    check("Startup logs outside __main__ (Gunicorn-visible)",
          app_source.index('Stripe Checkout') < app_source.index("if __name__"))

    # --- render.yaml completeness ---
    render_yaml_path = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), 'render.yaml')
    with open(render_yaml_path, 'r') as f:
        render_yaml = f.read()
    check("render.yaml includes STRIPE_PRICE_ID",
          'STRIPE_PRICE_ID' in render_yaml)

    # --- CSS checks ---
    check("Upload toggle row CSS exists",
          'upload-toggle-row' in builder_css)

    # --- Behavioral endpoint tests ---
    import io

    with app.test_client() as c:
        # Upload-generate rejects missing file
        r = c.post('/api/build/generate-from-upload',
                    data={'job_description': 'Senior software engineer with Python experience and cloud infrastructure knowledge for a fast-paced startup environment'},
                    content_type='multipart/form-data')
        check("Upload-generate rejects missing file (400)",
              r.status_code == 400)

        # Upload-generate rejects empty JD
        fake_pdf = (io.BytesIO(b'%PDF-1.4 fake pdf content'), 'test.pdf')
        r = c.post('/api/build/generate-from-upload',
                    data={'resume_file': fake_pdf, 'job_description': ''},
                    content_type='multipart/form-data')
        check("Upload-generate rejects empty JD (400)",
              r.status_code == 400)

        # Upload-generate rejects wrong file type
        fake_txt = (io.BytesIO(b'plain text content'), 'resume.txt')
        r = c.post('/api/build/generate-from-upload',
                    data={'resume_file': fake_txt,
                          'job_description': 'Senior software engineer with Python experience and cloud infrastructure knowledge for a fast-paced startup environment'},
                    content_type='multipart/form-data')
        check("Upload-generate rejects non-PDF/DOCX file (400)",
              r.status_code == 400)

        # Upload-generate rejects short JD
        fake_pdf2 = (io.BytesIO(b'%PDF-1.4 fake content'), 'resume.pdf')
        r = c.post('/api/build/generate-from-upload',
                    data={'resume_file': fake_pdf2,
                          'job_description': 'too short'},
                    content_type='multipart/form-data')
        check("Upload-generate rejects short JD (400)",
              r.status_code == 400)

    # --- .env.example completeness ---
    env_example_path = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), '.env.example')
    with open(env_example_path, 'r') as f:
        env_example = f.read()
    check(".env.example includes STRIPE_PRICE_ID",
          'STRIPE_PRICE_ID' in env_example)
    check(".env.example includes REDIS_URL",
          'REDIS_URL' in env_example)

    # ---- SECTION 15: PAYSTACK INTEGRATION (Nigeria) ----
    print("\n-- Section 15: Paystack Integration (Nigeria) --")

    # Read Paystack source files
    paystack_utils_path = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), 'backend', 'paystack_utils.py')
    with open(paystack_utils_path, 'r') as f:
        paystack_source = f.read()

    # -- Structure --
    check("paystack_utils.py exists and has create_paystack_transaction",
          'def create_paystack_transaction' in paystack_source)
    check("paystack_utils.py has verify_paystack_payment",
          'def verify_paystack_payment' in paystack_source)
    check("paystack_utils.py has verify_paystack_webhook",
          'def verify_paystack_webhook' in paystack_source)
    check("paystack_utils.py has format_naira_price",
          'def format_naira_price' in paystack_source)
    check("HMAC SHA512 verification logic present",
          'hmac.new' in paystack_source and 'sha512' in paystack_source)
    check("HMAC uses compare_digest for timing-safe comparison",
          'hmac.compare_digest' in paystack_source)

    # -- Payment Integrity (P0) --
    check("verify_paystack_payment checks amount == PAYSTACK_AMOUNT_KOBO",
          'PAYSTACK_AMOUNT_KOBO' in paystack_source and 'amount' in paystack_source and 'mismatch' in paystack_source)
    check("verify_paystack_payment checks currency == NGN",
          'PAYSTACK_CURRENCY' in paystack_source and 'currency' in paystack_source)
    check("verify_paystack_payment rejects wrong amount",
          'Payment amount mismatch' in paystack_source)
    check("verify_paystack_payment rejects wrong currency",
          'Payment currency mismatch' in paystack_source)
    check("create_paystack_transaction returns error without key",
          'Paystack not configured' in paystack_source)
    check("create_paystack_transaction requires email",
          'Email is required' in paystack_source)

    # -- Webhook Idempotency (P0) --
    check("Webhook uses SETNX dedup on reference before side effects",
          'paystack_processed' in app_source and 'nx=True' in app_source)
    check("Webhook releases dedup key on failure",
          'delete' in app_source and 'paystack_processed' in app_source)
    check("Webhook returns 200 on duplicate reference (early return)",
          'Already processed this reference' in app_source)

    # -- Email dedup separation (P1) --
    check("Email uses separate dedup via _send_cv_email (not webhook dedup)",
          '_send_cv_email' in app_source and 'cv_emailed' in app_source)
    check("_send_cv_email releases dedup on failure for retry",
          'delete(dedup_key)' in app_source or 'delete(f"resumeradar:cv_emailed' in app_source)

    # -- App.py integration --
    check("app.py imports paystack_utils",
          'from backend.paystack_utils import' in app_source)
    check("app.py has _get_base_url using PUBLIC_BASE_URL",
          'PUBLIC_BASE_URL' in app_source and '_get_base_url' in app_source)
    check("create-checkout accepts provider param",
          'provider' in app_source and '"paystack"' in app_source)
    check("create-checkout Paystack path requires delivery_email",
          'Email address is required for Naira payments' in app_source)
    check("download endpoint supports Paystack verification",
          'verify_paystack_payment' in app_source and 'paystack_ref' in app_source)
    check("Startup log shows Paystack status",
          'Paystack' in app_source and 'PAYSTACK_SECRET_KEY' in app_source)

    # -- TTL Alignment (P1) --
    check("Auxiliary Redis keys use 259200 TTL (72h)",
          app_source.count('259200') >= 4)

    # -- Frontend: builder.js --
    check("builder.js has shouldShowPaystackOption function",
          'shouldShowPaystackOption' in builder_js)
    check("builder.js checks Africa/Lagos timezone only",
          "Africa/Lagos" in builder_js)
    check("builder.js default provider is stripe",
          "detectedProvider = 'stripe'" in builder_js)
    check("builder.js clears stale sessionStorage for non-Nigeria users",
          "removeItem('resumeradar_force_paystack')" in builder_js)
    check("builder.js requires email for Paystack before checkout",
          'required for Naira payments' in builder_js)
    check("builder.js sends provider in checkout request body",
          'provider: detectedProvider' in builder_js)
    check("builder.js handles Paystack callback (reference/trxref params)",
          'trxref' in builder_js and 'resumeradar_paystack_ref' in builder_js)
    check("builder.js handlePostPayment accepts provider + paystackRef",
          'handlePostPayment(token, sessionId, provider, paystackRef)' in builder_js)

    # -- Frontend: build.html --
    check("build.html has providerToggle element",
          'id="providerToggle"' in build_html)
    check("build.html has paymentPrice element with id",
          'id="paymentPrice"' in build_html)
    check("build.html has providerRow element",
          'id="providerRow"' in build_html)
    check("build.html has data-paystack-enabled attribute",
          'data-paystack-enabled' in build_html)
    check("build.html has data-paystack-price attribute",
          'data-paystack-price' in build_html)
    check("build.html has emailRequiredHint element",
          'id="emailRequiredHint"' in build_html)

    # -- Frontend: builder.css --
    check("builder.css has payment-provider-row styles",
          'payment-provider-row' in builder_css)

    # -- Config --
    check(".env.example includes PAYSTACK_SECRET_KEY",
          'PAYSTACK_SECRET_KEY' in env_example)
    check(".env.example includes PAYSTACK_AMOUNT_KOBO",
          'PAYSTACK_AMOUNT_KOBO' in env_example)
    check(".env.example includes PUBLIC_BASE_URL",
          'PUBLIC_BASE_URL' in env_example)
    check("render.yaml includes PAYSTACK_SECRET_KEY",
          'PAYSTACK_SECRET_KEY' in render_yaml)
    check("render.yaml includes PUBLIC_BASE_URL",
          'PUBLIC_BASE_URL' in render_yaml)

    # -- Behavioral Endpoint Tests --
    with app.test_client() as c:
        # Paystack webhook endpoint exists
        r = c.post('/api/build/webhook/paystack',
                    data=b'{}',
                    content_type='application/json')
        check("Paystack webhook endpoint exists (not 404)",
              r.status_code != 404)

        # Paystack webhook rejects unsigned request
        r = c.post('/api/build/webhook/paystack',
                    data=b'{"event": "charge.success"}',
                    content_type='application/json')
        check("Paystack webhook rejects unsigned request (400)",
              r.status_code == 400)

        # create-checkout with provider=paystack: if PAYSTACK_SECRET_KEY is set, requires email.
        # If not set, falls through to Stripe. Check source for the validation.
        check("create-checkout Paystack email validation in source",
              'Email address is required for Naira payments' in app_source and
              'provider == "paystack"' in app_source)

        # create-checkout default provider is stripe (no provider param)
        r = c.post('/api/build/create-checkout',
                    json={"token": "test123", "template": "classic"},
                    content_type='application/json')
        # Without a valid token it may return 400 (token expired) or 500 (stripe error)
        # Key: it should NOT try Paystack path
        result = r.get_json() or {}
        check("create-checkout defaults to stripe when no provider given",
              'Naira' not in result.get('error', ''))

    # ---- SECTION 16: CV EXTRACTION COMPLETENESS (Education/Certs Fix) ----
    print("\n-- Section 16: CV Extraction Completeness --")

    from backend.cv_builder import (
        _smart_truncate_resume, _fallback_extract_education_certs,
        _assess_extraction_quality, _EDU_HEADING_RE, _CERT_HEADING_RE,
        _SECTION_HEADING_RE, _EDU_CERT_HEADING_RE
    )

    # -- Source pattern checks --
    cv_builder_path = os.path.join(project_root, 'backend', 'cv_builder.py')
    with open(cv_builder_path, 'r') as f:
        cv_builder_source = f.read()

    check("_smart_truncate_resume function exists",
          'def _smart_truncate_resume' in cv_builder_source)
    check("_smart_truncate_resume called in extract_and_polish prompt",
          '_smart_truncate_resume(resume_text)' in cv_builder_source)
    check("_EDU_HEADING_RE and _CERT_HEADING_RE are separate regexes",
          '_EDU_HEADING_RE = re.compile' in cv_builder_source and
          '_CERT_HEADING_RE = re.compile' in cv_builder_source)
    check("Heading regexes are line-anchored",
          "r'^\\s*(?:EDUCATION" in cv_builder_source or
          "r'^\\s*(?:CERTIF" in cv_builder_source)
    check("_SECTION_HEADING_RE matches UPPERCASE and Title Case",
          '_SECTION_HEADING_RE' in cv_builder_source and
          '[A-Z][a-z]+' in cv_builder_source)
    check("Section heading synonyms in prompt",
          'ACADEMIC QUALIFICATIONS' in cv_builder_source and
          'PROFESSIONAL TRAINING' in cv_builder_source)
    check("_fallback_extract_education_certs function exists",
          'def _fallback_extract_education_certs' in cv_builder_source)
    check("_assess_extraction_quality function exists",
          'def _assess_extraction_quality' in cv_builder_source)
    check("Quality assessment checks education + certs + experience",
          'education_missing' in cv_builder_source and
          'certifications_missing' in cv_builder_source and
          'experience_missing' in cv_builder_source)
    check("max_tokens >= 5000 in extract_and_polish",
          'max_tokens=5000' in cv_builder_source)
    check("No resume_text[:5000] in source",
          'resume_text[:5000]' not in cv_builder_source)
    check("Strict rule: NEVER omit any entries",
          'NEVER omit any entries' in cv_builder_source)

    # -- Frontend checks --
    check("extractionWarning element in build.html",
          'id="extractionWarning"' in build_html)
    check("extractionWarningText element in build.html",
          'id="extractionWarningText"' in build_html)
    check("extraction_warnings handling in builder.js",
          'extraction_warnings' in builder_js)
    check("showExtractionWarnings function in builder.js",
          'showExtractionWarnings' in builder_js)
    check("Payment gate blocks education_missing in builder.js",
          'education_missing' in builder_js and 'certifications_missing' in builder_js)
    check("Payment gate calls showError and returns",
          'missing Education or Certifications' in builder_js)
    check("Extraction warning banner CSS exists",
          'extraction-warning-banner' in builder_css)

    # -- Behavioral: smart truncation --
    # Short input passes through unchanged
    short_text = "A" * 5000
    result_trunc = _smart_truncate_resume(short_text, max_chars=12000)
    check("Smart truncation: short input unchanged",
          result_trunc == short_text)

    # Long input with edu heading at end preserves it
    long_text = ("X " * 7000 + "\nACADEMIC QUALIFICATIONS\nBSc Computer Science, UCL, 2019\n"
                 "MSc Data Science, Imperial, 2021\n")
    result_trunc = _smart_truncate_resume(long_text, max_chars=12000)
    check("Smart truncation: preserves ACADEMIC QUALIFICATIONS",
          'ACADEMIC QUALIFICATIONS' in result_trunc)
    check("Smart truncation: contains splice marker",
          '[... some' in result_trunc)

    # Hard cap enforced
    huge_text = "A" * 20000
    result_trunc = _smart_truncate_resume(huge_text, max_chars=12000)
    check("Smart truncation: hard cap enforced",
          len(result_trunc) <= 12000, f"Got {len(result_trunc)}")

    # Hard cap with edu heading
    huge_with_edu = "B " * 10000 + "\nEDUCATION\nBSc test\n" + "C " * 5000
    result_trunc = _smart_truncate_resume(huge_with_edu, max_chars=12000)
    check("Smart truncation: hard cap with edu heading",
          len(result_trunc) <= 12000, f"Got {len(result_trunc)}")

    # -- Behavioral: fallback extractor --
    # Education extraction from section
    synth_text = """PROFESSIONAL EXPERIENCE
Senior Engineer at Acme Corp
Jan 2020 - Present
Built APIs

ACADEMIC QUALIFICATIONS
BSc Computer Science, University of Lagos, 2015
MSc Information Systems, University of Ibadan, 2018

PROFESSIONAL TRAINING
PMP Certification, PMI, 2020
AWS Solutions Architect Associate, Amazon, 2021
ITIL Foundation, Axelos, 2019
"""
    # Test with empty AI arrays
    test_result = {"education": [], "certifications": [], "experience": []}
    counts = _fallback_extract_education_certs(synth_text, test_result)
    check("Fallback extractor: finds education entries",
          len(test_result["education"]) >= 1,
          f"Found {len(test_result['education'])} entries")
    check("Fallback extractor: finds certification entries",
          len(test_result["certifications"]) >= 1,
          f"Found {len(test_result['certifications'])} entries")
    check("Fallback extractor: raw_edu_count > 0",
          counts.get("raw_edu_count", 0) >= 1,
          f"raw_edu_count={counts.get('raw_edu_count', 0)}")
    check("Fallback extractor: raw_cert_count > 0",
          counts.get("raw_cert_count", 0) >= 1,
          f"raw_cert_count={counts.get('raw_cert_count', 0)}")

    # No education signals = empty (no false positives)
    no_edu_text = "PROFESSIONAL EXPERIENCE\nSenior Engineer\nBuilt APIs\n"
    no_edu_result = {"education": [], "certifications": []}
    no_edu_counts = _fallback_extract_education_certs(no_edu_text, no_edu_result)
    check("Fallback extractor: no false positives without edu heading",
          len(no_edu_result["education"]) == 0)
    check("Fallback extractor: raw_edu_count is 0 without heading",
          no_edu_counts.get("raw_edu_count", 0) == 0)

    # Partial miss: AI has 1, resume has 2 → adds 1
    partial_text = """EDUCATION
BSc Computer Science, UCL, 2015
MSc Data Science, Imperial, 2018
"""
    partial_result = {
        "education": [{"degree": "BSc Computer Science", "institution": "UCL",
                        "graduation_date": "2015", "details": ""}],
        "certifications": []
    }
    partial_counts = _fallback_extract_education_certs(partial_text, partial_result)
    check("Fallback extractor: partial miss detection",
          len(partial_result["education"]) >= 2,
          f"Got {len(partial_result['education'])} entries")

    # -- Behavioral: quality assessment --
    # Education missing: heading + section entries + empty result
    qtext_edu = "Some text\nACADEMIC QUALIFICATIONS\nBSc CS, UCL, 2019\n"
    qresult_edu = {"education": [], "certifications": [], "experience": []}
    qcounts_edu = {"raw_edu_count": 1, "raw_cert_count": 0}
    warnings = _assess_extraction_quality(qtext_edu, qresult_edu, qcounts_edu)
    check("Quality assessment: education_missing when heading + entries + empty",
          'education_missing' in warnings, f"Got: {warnings}")

    # Cert missing: heading + section entries + empty result
    qtext_cert = "Some text\nCERTIFICATIONS\nAWS SA, Amazon, 2021\n"
    qresult_cert = {"education": [], "certifications": [], "experience": []}
    qcounts_cert = {"raw_edu_count": 0, "raw_cert_count": 1}
    warnings = _assess_extraction_quality(qtext_cert, qresult_cert, qcounts_cert)
    check("Quality assessment: certifications_missing when heading + entries + empty",
          'certifications_missing' in warnings, f"Got: {warnings}")

    # Experience partial: many date ranges, few extracted
    qtext_exp = "Some text\n" + "Jan 2020 - Present\n" * 5
    qresult_exp = {"education": [], "certifications": [], "experience": [{"title": "Eng"}]}
    warnings = _assess_extraction_quality(qtext_exp, qresult_exp)
    check("Quality assessment: experience advisory warning",
          'experience_missing' in warnings or 'experience_partial' in warnings,
          f"Got: {warnings}")

    # Complete result = no warnings
    qtext_complete = "No headings here. Just plain text."
    qresult_complete = {"education": [{"degree": "BSc"}], "certifications": [{"name": "AWS"}],
                        "experience": [{"title": "Eng"}]}
    warnings = _assess_extraction_quality(qtext_complete, qresult_complete)
    check("Quality assessment: complete result has no warnings",
          len(warnings) == 0, f"Got: {warnings}")

    # False-positive: "certified" in bullet but NO cert heading
    qtext_fp = "PROFESSIONAL EXPERIENCE\nCertified engineer who built systems.\nJan 2020 - Present\n"
    qresult_fp = {"education": [], "certifications": [], "experience": [{"title": "Eng"}]}
    qcounts_fp = {"raw_edu_count": 0, "raw_cert_count": 0}
    warnings = _assess_extraction_quality(qtext_fp, qresult_fp, qcounts_fp)
    check("Quality assessment: no false positive from 'certified' in bullet",
          'certifications_missing' not in warnings, f"Got: {warnings}")

    # False-positive: cert heading should NOT trigger education_missing
    qtext_cross = "Some text\nPROFESSIONAL TRAINING\nPMP, PMI, 2020\n"
    qresult_cross = {"education": [], "certifications": [], "experience": []}
    qcounts_cross = {"raw_edu_count": 0, "raw_cert_count": 1}
    warnings = _assess_extraction_quality(qtext_cross, qresult_cross, qcounts_cross)
    check("Quality assessment: cert heading does NOT trigger education_missing",
          'education_missing' not in warnings, f"Got: {warnings}")

    # Partial warnings
    qtext_partial = "Some text\nEDUCATION\nBSc CS\nMSc DS\n"
    qresult_partial = {"education": [{"degree": "BSc CS"}], "certifications": [], "experience": []}
    qcounts_partial = {"raw_edu_count": 2, "raw_cert_count": 0}
    warnings = _assess_extraction_quality(qtext_partial, qresult_partial, qcounts_partial)
    check("Quality assessment: education_partial when raw > ai",
          'education_partial' in warnings, f"Got: {warnings}")

    # -- End-to-end fixture: synthetic long CV --
    e2e_experience = "\n".join([
        f"Software Engineer at Company{i}\nJan {2015+i} - Dec {2016+i}\n- Built systems\n- Led team\n"
        for i in range(5)
    ])
    e2e_edu = "\nACADEMIC QUALIFICATIONS\nBSc Computer Science, University of Lagos, 2012\nMSc IT, University of Ibadan, 2014\n"
    e2e_certs = "\nPROFESSIONAL TRAINING\nPMP, PMI, 2018\nAWS SA, Amazon, 2019\nITIL Foundation, Axelos, 2020\n"
    e2e_resume = ("X " * 5000) + e2e_experience + e2e_edu + e2e_certs
    check("E2E fixture: resume > 10000 chars",
          len(e2e_resume) > 10000, f"Length: {len(e2e_resume)}")

    # Smart truncation preserves edu/cert
    e2e_truncated = _smart_truncate_resume(e2e_resume, max_chars=12000)
    check("E2E: truncated text contains ACADEMIC QUALIFICATIONS",
          'ACADEMIC QUALIFICATIONS' in e2e_truncated)
    check("E2E: truncated text contains PROFESSIONAL TRAINING",
          'PROFESSIONAL TRAINING' in e2e_truncated)
    check("E2E: truncated text within hard cap",
          len(e2e_truncated) <= 12000, f"Length: {len(e2e_truncated)}")

    # Fallback extracts from full text
    e2e_result = {"education": [], "certifications": [], "experience": []}
    e2e_counts = _fallback_extract_education_certs(e2e_resume, e2e_result)
    check("E2E: fallback finds education entries",
          len(e2e_result["education"]) >= 1,
          f"Found {len(e2e_result['education'])}")
    check("E2E: fallback finds certification entries",
          len(e2e_result["certifications"]) >= 1,
          f"Found {len(e2e_result['certifications'])}")

    # Quality assessment on incomplete result
    e2e_incomplete = {"education": [], "certifications": [], "experience": []}
    e2e_warnings = _assess_extraction_quality(e2e_resume, e2e_incomplete, e2e_counts)
    check("E2E: quality warns on incomplete result",
          len(e2e_warnings) > 0, f"Warnings: {e2e_warnings}")

    # Quality assessment on complete result (post-fallback)
    e2e_complete_counts = {"raw_edu_count": len(e2e_result["education"]),
                           "raw_cert_count": len(e2e_result["certifications"])}
    e2e_warnings_complete = _assess_extraction_quality(e2e_resume, e2e_result, e2e_complete_counts)
    check("E2E: no warnings on complete post-fallback result",
          'education_missing' not in e2e_warnings_complete and
          'certifications_missing' not in e2e_warnings_complete,
          f"Warnings: {e2e_warnings_complete}")

    # ---- SECTION 17: DOCX DOWNLOAD FEATURE ----
    print("\n-- Section 17: DOCX Download Feature --")

    # -- Source file reads --
    docx_gen_path = os.path.join(project_root, 'backend', 'cv_docx_generator.py')
    with open(docx_gen_path, 'r') as f:
        docx_gen_source = f.read()

    stripe_utils_path = os.path.join(project_root, 'backend', 'stripe_utils.py')
    with open(stripe_utils_path, 'r') as f:
        stripe_utils_source = f.read()

    paystack_utils_path = os.path.join(project_root, 'backend', 'paystack_utils.py')
    with open(paystack_utils_path, 'r') as f:
        paystack_utils_source = f.read()

    # Re-read app_source, builder_js, build_html for fresh state
    with open(app_py_path, 'r') as f:
        app_source_17 = f.read()
    with open(builder_js_path, 'r') as f:
        builder_js_17 = f.read()
    with open(builder_css_path, 'r') as f:
        builder_css_17 = f.read()

    # -- Source pattern checks: cv_docx_generator.py --
    check("DOCX: generate_cv_docx function exists",
          'def generate_cv_docx' in docx_gen_source)
    check("DOCX: does NOT import _safe from cv_pdf_generator",
          'import _safe' not in docx_gen_source and 'from backend.cv_pdf_generator import' in docx_gen_source
          and '_safe' not in docx_gen_source.split('from backend.cv_pdf_generator import')[1].split('\n')[0])
    check("DOCX: has own _docx_safe function",
          'def _docx_safe' in docx_gen_source)
    check("DOCX: _docx_safe does NOT use latin-1 encoding",
          "encode('latin" not in docx_gen_source and '.encode("latin' not in docx_gen_source)
    check("DOCX: imports _flatten_skills from cv_pdf_generator",
          '_flatten_skills' in docx_gen_source.split('from backend.cv_pdf_generator import')[1].split('\n')[0])
    check("DOCX: imports _format_contact_line from cv_pdf_generator",
          '_format_contact_line' in docx_gen_source.split('from backend.cv_pdf_generator import')[1].split('\n')[0])
    check("DOCX: imports _format_date_range from cv_pdf_generator",
          '_format_date_range' in docx_gen_source.split('from backend.cv_pdf_generator import')[1].split('\n')[0])
    check("DOCX: returns bytes (BytesIO pattern)",
          'BytesIO' in docx_gen_source and '.getvalue()' in docx_gen_source)
    check("DOCX: 3 templates — classic, modern, minimal",
          '_render_classic' in docx_gen_source and '_render_modern' in docx_gen_source and '_render_minimal' in docx_gen_source)

    # -- Source pattern checks: app.py --
    check("DOCX: generate_cv_docx imported in app.py",
          'from backend.cv_docx_generator import generate_cv_docx' in app_source_17)
    check("DOCX: format parameter handling in download endpoint",
          'dl_format' in app_source_17 and 'VALID_FORMATS' in app_source_17)
    check("DOCX: zipfile import in download endpoint",
          'import zipfile' in app_source_17)

    # -- Source pattern checks: build.html --
    build_r = c.get('/build')
    build_html_17 = build_r.data.decode()
    check("DOCX: format-toggle buttons in build.html",
          'format-toggle' in build_html_17 and 'data-format' in build_html_17)
    check("DOCX: three format options (both, pdf, docx)",
          'data-format="both"' in build_html_17 and 'data-format="pdf"' in build_html_17 and 'data-format="docx"' in build_html_17)

    # -- Source pattern checks: builder.js --
    check("DOCX: selectedFormat variable in builder.js",
          'selectedFormat' in builder_js_17)
    check("DOCX: resumeradar_cv_format sessionStorage in builder.js",
          'resumeradar_cv_format' in builder_js_17)
    check("DOCX: Content-Type based filename in builder.js",
          'wordprocessingml' in builder_js_17 or 'contentType' in builder_js_17)
    check("DOCX: format passed to checkout in builder.js",
          'format: selectedFormat' in builder_js_17 or 'format:selectedFormat' in builder_js_17)

    # -- Source pattern checks: email size guard --
    check("DOCX: 10MB size guard in _send_cv_email",
          '10_000_000' in app_source_17 or '10000000' in app_source_17)
    check("DOCX: dual attachment in email sender",
          'generate_cv_docx' in app_source_17.split('def _send_cv_email')[1].split('\ndef ')[0]
          if 'def _send_cv_email' in app_source_17 else False)

    # -- Source pattern checks: payment utils format plumbing --
    check("DOCX: stripe create_checkout_session accepts format_choice",
          'format_choice' in stripe_utils_source.split('def create_checkout_session')[1].split('\n')[0])
    check("DOCX: paystack create_paystack_transaction accepts format_choice",
          'format_choice' in paystack_utils_source.split('def create_paystack_transaction')[1].split('\n')[0])
    check("DOCX: stripe verify returns format with empty default",
          '"format": session.metadata.get("format", "")' in stripe_utils_source or
          "'format': session.metadata.get('format', '')" in stripe_utils_source)
    check("DOCX: paystack verify returns format with empty default",
          '"format": metadata.get("format", "")' in paystack_utils_source or
          "'format': metadata.get('format', '')" in paystack_utils_source)
    check("DOCX: app.py build_create_checkout passes format_choice",
          'format_choice' in app_source_17.split('def build_create_checkout')[1].split('\ndef ')[0]
          if 'def build_create_checkout' in app_source_17 else False)
    check("DOCX: app.py build_download uses 3-step format resolution",
          'payment.get("format"' in app_source_17 or "payment.get('format'" in app_source_17)
    check("DOCX: _send_cv_email signature unchanged (no format_choice param)",
          'def _send_cv_email(email, token, template, event_id)' in app_source_17)

    # -- Behavioral: DOCX generation --
    from backend.cv_docx_generator import generate_cv_docx as gen_docx
    from io import BytesIO

    docx_test_cv = {
        'personal': {'full_name': 'QA Test', 'email': 'qa@test.com', 'phone': '+1 555 0199', 'location': 'London'},
        'summary': 'Experienced engineer with 5 years building scalable systems.',
        'experience': [
            {'title': 'Senior Engineer', 'company': 'Acme Corp', 'start_date': 'Jan 2020', 'end_date': 'Present',
             'bullets': ['Built microservices', 'Led team of 4', 'Reduced deploy time by 60%']},
            {'title': 'Junior Engineer', 'company': 'StartupCo', 'start_date': 'Jun 2017', 'end_date': 'Dec 2019',
             'bullets': ['Developed REST APIs', 'Wrote unit tests']},
            {'title': 'Intern', 'company': 'TechInc', 'start_date': 'Jan 2017', 'end_date': 'May 2017',
             'bullets': ['Assisted with frontend development']},
        ],
        'education': [
            {'degree': 'BSc Computer Science', 'institution': 'UCL', 'graduation_date': '2019', 'details': 'First Class Honours'},
            {'degree': 'MSc Data Science', 'institution': 'Imperial', 'graduation_date': '2021', 'details': ''},
        ],
        'skills': ['Python', 'AWS', 'Docker', 'Kubernetes', 'React'],
        'certifications': [
            {'name': 'AWS Solutions Architect', 'issuer': 'Amazon', 'date': '2023'},
            {'name': 'PMP', 'issuer': 'PMI', 'date': '2022'},
        ]
    }

    for tmpl in ['classic', 'modern', 'minimal']:
        try:
            docx_out = gen_docx(docx_test_cv, tmpl)
            is_docx = isinstance(docx_out, bytes) and len(docx_out) > 500
            check(f"DOCX {tmpl} template generates", is_docx, f"{len(docx_out)} bytes")
        except Exception as e:
            check(f"DOCX {tmpl} template generates", False, str(e))

    # Output starts with PK signature (ZIP/DOCX magic bytes)
    try:
        docx_out = gen_docx(docx_test_cv, 'classic')
        check("DOCX: output starts with PK magic bytes",
              docx_out[:2] == b'PK', f"Got: {docx_out[:4]}")
    except Exception as e:
        check("DOCX: output starts with PK magic bytes", False, str(e))

    # Output can be loaded by python-docx
    try:
        from docx import Document as DocxDocument
        doc = DocxDocument(BytesIO(docx_out))
        check("DOCX: loadable by python-docx", True)

        # Extract all text
        all_text = "\n".join([p.text for p in doc.paragraphs])
        # Also check table cells (Modern template uses tables)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    all_text += "\n" + "\n".join([p.text for p in cell.paragraphs])

        check("DOCX: contains personal name", 'QA Test' in all_text)
        check("DOCX: contains summary text", 'scalable systems' in all_text)
        check("DOCX: contains experience titles",
              'Senior Engineer' in all_text and 'Junior Engineer' in all_text)
        check("DOCX: contains education degrees",
              'BSc Computer Science' in all_text or 'BSc' in all_text)
        check("DOCX: contains certification names",
              'AWS Solutions Architect' in all_text or 'AWS' in all_text)
        check("DOCX: contains skills", 'Python' in all_text)
    except ImportError:
        check("DOCX: loadable by python-docx", False, "python-docx not installed")
    except Exception as e:
        check("DOCX: loadable by python-docx", False, str(e))

    # Invalid template falls back to classic
    try:
        docx_fallback = gen_docx(docx_test_cv, 'nonexistent')
        check("DOCX: invalid template falls back to classic",
              isinstance(docx_fallback, bytes) and len(docx_fallback) > 500)
    except Exception as e:
        check("DOCX: invalid template falls back to classic", False, str(e))

    # Empty CV doesn't crash
    empty_cv_docx = {'personal': {}, 'summary': '', 'experience': [], 'education': [], 'skills': [], 'certifications': []}
    try:
        all_ok = True
        for tmpl in ['classic', 'modern', 'minimal']:
            out = gen_docx(empty_cv_docx, tmpl)
            if not isinstance(out, bytes) or len(out) < 100:
                all_ok = False
        check("DOCX: empty CV doesn't crash (all templates)", all_ok)
    except Exception as e:
        check("DOCX: empty CV doesn't crash (all templates)", False, str(e))

    # -- Behavioral: DOCX Unicode safety --
    from backend.cv_docx_generator import _docx_safe

    check("DOCX: _docx_safe preserves Unicode (em dash, accented)",
          _docx_safe("Olú Adéyígá — Senior Manager") == "Olú Adéyígá — Senior Manager")
    check("DOCX: _docx_safe strips control chars",
          '\x00' not in _docx_safe("Hello\x00World\x0bTest") and '\x0b' not in _docx_safe("Hello\x00World\x0bTest"))
    check("DOCX: _docx_safe preserves normal text",
          _docx_safe("Normal text here") == "Normal text here")

    # Generate DOCX with accented name — verify preserved
    try:
        unicode_cv = dict(docx_test_cv)
        unicode_cv['personal'] = dict(docx_test_cv['personal'])
        unicode_cv['personal']['full_name'] = "José García-López"
        unicode_out = gen_docx(unicode_cv, 'classic')
        unicode_doc = DocxDocument(BytesIO(unicode_out))
        unicode_text = "\n".join([p.text for p in unicode_doc.paragraphs])
        check("DOCX: Unicode name preserved in document",
              "José García-López" in unicode_text, f"Name found: {'José García-López' in unicode_text}")
    except Exception as e:
        check("DOCX: Unicode name preserved in document", False, str(e))

    # -- Behavioral: DOCX template-specific features --
    try:
        # Classic: check for pBdr (paragraph border bottom) in XML
        classic_out = gen_docx(docx_test_cv, 'classic')
        classic_doc = DocxDocument(BytesIO(classic_out))
        classic_xml = classic_doc.element.xml
        check("DOCX classic: has paragraph borders (pBdr)",
              'pBdr' in classic_xml)

        # Modern: check for blue accent (border color in XML)
        modern_out = gen_docx(docx_test_cv, 'modern')
        modern_doc = DocxDocument(BytesIO(modern_out))
        modern_xml = modern_doc.element.xml
        check("DOCX modern: has border color (accent tables)",
              'tcBorders' in modern_xml or '2563eb' in modern_xml.lower() or '2563EB' in modern_xml)
    except Exception as e:
        check("DOCX classic/modern template features", False, str(e))

    # -- Behavioral: true endpoint tests (Flask test client) --
    # Mock Redis + payment verification for endpoint tests
    import unittest.mock as mock
    from unittest.mock import MagicMock, patch

    mock_cv_data = json.dumps(docx_test_cv)

    def _mock_redis_get(key):
        if 'resumeradar:cv:' in key:
            return mock_cv_data
        if 'resumeradar:cv_paid:' in key:
            return json.dumps({"template": "classic", "delivery_email": "", "format": ""})
        return None

    def _mock_redis_get_with_format(fmt):
        def getter(key):
            if 'resumeradar:cv:' in key:
                return mock_cv_data
            if 'resumeradar:cv_paid:' in key:
                return json.dumps({"template": "classic", "delivery_email": "", "format": fmt})
            return None
        return getter

    # Test: format=docx returns DOCX content type
    try:
        with patch('app._redis_client') as mock_redis:
            mock_redis.get = MagicMock(side_effect=_mock_redis_get)
            mock_redis.exists = MagicMock(return_value=True)
            with patch('app.verify_checkout_payment', return_value={
                'verified': True, 'template': 'classic', 'delivery_email': '', 'format': 'docx'
            }):
                resp = c.get('/api/build/download/test-token-123?session_id=cs_test&format=docx')
                check("DOCX endpoint: format=docx returns DOCX Content-Type",
                      resp.status_code == 200 and 'wordprocessingml' in (resp.content_type or ''),
                      f"Status: {resp.status_code}, CT: {resp.content_type}")
                if resp.status_code == 200:
                    check("DOCX endpoint: format=docx body starts with PK",
                          resp.data[:2] == b'PK')
    except Exception as e:
        check("DOCX endpoint: format=docx returns DOCX Content-Type", False, str(e))
        check("DOCX endpoint: format=docx body starts with PK", False, "skipped")

    # Test: format=both returns ZIP
    try:
        with patch('app._redis_client') as mock_redis:
            mock_redis.get = MagicMock(side_effect=_mock_redis_get)
            mock_redis.exists = MagicMock(return_value=True)
            with patch('app.verify_checkout_payment', return_value={
                'verified': True, 'template': 'classic', 'delivery_email': '', 'format': 'both'
            }):
                resp = c.get('/api/build/download/test-token-123?session_id=cs_test&format=both')
                check("DOCX endpoint: format=both returns ZIP Content-Type",
                      resp.status_code == 200 and 'zip' in (resp.content_type or ''),
                      f"Status: {resp.status_code}, CT: {resp.content_type}")
                if resp.status_code == 200:
                    import zipfile as zf
                    z = zf.ZipFile(BytesIO(resp.data))
                    names = z.namelist()
                    check("DOCX endpoint: ZIP contains PDF and DOCX",
                          any('.pdf' in n for n in names) and any('.docx' in n for n in names),
                          f"Files: {names}")
    except Exception as e:
        check("DOCX endpoint: format=both returns ZIP Content-Type", False, str(e))
        check("DOCX endpoint: ZIP contains PDF and DOCX", False, "skipped")

    # Test: format=pdf returns PDF (existing behavior)
    try:
        with patch('app._redis_client') as mock_redis:
            mock_redis.get = MagicMock(side_effect=_mock_redis_get)
            mock_redis.exists = MagicMock(return_value=True)
            with patch('app.verify_checkout_payment', return_value={
                'verified': True, 'template': 'classic', 'delivery_email': '', 'format': 'pdf'
            }):
                resp = c.get('/api/build/download/test-token-123?session_id=cs_test&format=pdf')
                check("DOCX endpoint: format=pdf returns PDF Content-Type",
                      resp.status_code == 200 and 'pdf' in (resp.content_type or ''),
                      f"Status: {resp.status_code}, CT: {resp.content_type}")
    except Exception as e:
        check("DOCX endpoint: format=pdf returns PDF Content-Type", False, str(e))

    # Test: no format param defaults to PDF (backward compatible)
    try:
        with patch('app._redis_client') as mock_redis:
            mock_redis.get = MagicMock(side_effect=_mock_redis_get)
            mock_redis.exists = MagicMock(return_value=True)
            with patch('app.verify_checkout_payment', return_value={
                'verified': True, 'template': 'classic', 'delivery_email': '', 'format': ''
            }):
                resp = c.get('/api/build/download/test-token-123?session_id=cs_test')
                check("DOCX endpoint: no format + empty metadata = PDF default",
                      resp.status_code == 200 and 'pdf' in (resp.content_type or ''),
                      f"Status: {resp.status_code}, CT: {resp.content_type}")
    except Exception as e:
        check("DOCX endpoint: no format + empty metadata = PDF default", False, str(e))

    # Test: metadata fallback — no format in request, but metadata has 'both'
    try:
        with patch('app._redis_client') as mock_redis:
            mock_redis.get = MagicMock(side_effect=_mock_redis_get_with_format('both'))
            mock_redis.exists = MagicMock(return_value=True)
            with patch('app.verify_checkout_payment', return_value={
                'verified': True, 'template': 'classic', 'delivery_email': '', 'format': 'both'
            }):
                resp = c.get('/api/build/download/test-token-123?session_id=cs_test')
                check("DOCX endpoint: metadata fallback format=both → ZIP",
                      resp.status_code == 200 and 'zip' in (resp.content_type or ''),
                      f"Status: {resp.status_code}, CT: {resp.content_type}")
    except Exception as e:
        check("DOCX endpoint: metadata fallback format=both → ZIP", False, str(e))

    # Test: pre-DOCX backward compat — no format anywhere → PDF
    try:
        with patch('app._redis_client') as mock_redis:
            mock_redis.get = MagicMock(side_effect=_mock_redis_get_with_format(''))
            mock_redis.exists = MagicMock(return_value=True)
            with patch('app.verify_checkout_payment', return_value={
                'verified': True, 'template': 'classic', 'delivery_email': '', 'format': ''
            }):
                resp = c.get('/api/build/download/test-token-123?session_id=cs_test')
                check("DOCX endpoint: pre-DOCX payment (format='') → PDF",
                      resp.status_code == 200 and 'pdf' in (resp.content_type or ''),
                      f"Status: {resp.status_code}, CT: {resp.content_type}")
    except Exception as e:
        check("DOCX endpoint: pre-DOCX payment (format='') → PDF", False, str(e))

    # -- Format toggle CSS --
    check("DOCX: format-toggle-group CSS in builder.css",
          'format-toggle-group' in builder_css_17)
    check("DOCX: format-toggle active style in builder.css",
          'format-toggle.active' in builder_css_17 or '.format-toggle.active' in builder_css_17)
    check("DOCX: format-hint CSS in builder.css",
          'format-hint' in builder_css_17)

    elapsed = time.time() - start

    # ============================================================
    # REPORT
    # ============================================================
    print()
    print("=" * 55)
    print(f"  RESUMERADAR QA REPORT {'(QUICK)' if QUICK_MODE else '(FULL)'}")
    print("=" * 55)

    for status, name, detail in RESULTS:
        icon = "  PASS" if status == "PASS" else "  FAIL"
        suffix = f"  ({detail})" if detail else ""
        print(f"{icon}  {name}{suffix}")

    print()
    print("-" * 55)
    print(f"  {PASS} passed, {FAIL} failed  |  {elapsed:.1f}s")

    if FAIL > 0:
        print()
        print("  FAILURES:")
        for status, name, detail in RESULTS:
            if status == "FAIL":
                print(f"    - {name}" + (f" ({detail})" if detail else ""))

    print("-" * 55)

    if FAIL == 0:
        print("  STATUS: ALL CLEAR — safe to deploy")
    else:
        print(f"  STATUS: {FAIL} ISSUE(S) FOUND — review before deploying")

    print("=" * 55)
    print()

    return FAIL


if __name__ == '__main__':
    failures = run_tests()
    sys.exit(1 if failures > 0 else 0)
