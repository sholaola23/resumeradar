"""
ResumeRadar -- CV DOCX Generator
Generates ATS-friendly resume Word documents in 3 templates: Classic, Modern, Minimal.
Uses python-docx, Calibri font, single-column layout. Mirrors cv_pdf_generator.py structure.

Key differences from PDF generator:
- Uses Calibri (Word default, universally available) instead of Helvetica
- Full Unicode support — no latin-1 encoding (preserves accented names, em dashes, etc.)
- A4 page with 1-inch margins (Word default)
- Tab stops for right-aligned dates derived from text width calculation

Design matches the same 4 reference CVs and ATS best practices as the PDF version.
"""

import re
from io import BytesIO
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml, OxmlElement

from backend.cv_pdf_generator import _flatten_skills, _format_contact_line, _format_date_range


# ============================================================
# CONSTANTS
# ============================================================

TEMPLATES = ["classic", "modern", "minimal"]

# A4 page: 21cm width, 2.54cm margins each side → text width = 15.92cm
# This single canonical value is used for all right-aligned tab stops.
TEXT_WIDTH_CM = 15.92

# Color palette (matching PDF generator)
COLOR_DARK = RGBColor(31, 41, 55)       # #1f2937 — headings, names
COLOR_BODY = RGBColor(55, 65, 81)       # #374151 — body text
COLOR_GRAY = RGBColor(107, 114, 128)    # #6b7280 — dates, contact info
COLOR_BLUE = RGBColor(37, 99, 235)      # #2563eb — modern template accent


# ============================================================
# TEXT SANITIZER (Unicode-safe — NOT the PDF _safe() function)
# ============================================================

# Control characters that are invalid in DOCX XML (except \t, \n, \r)
_CONTROL_CHAR_RE = re.compile(r'[\x00-\x08\x0b\x0c\x0e-\x1f]')


def _docx_safe(text):
    """Sanitize text for DOCX output.

    Unlike the PDF _safe() which forces latin-1 encoding (degrading Unicode),
    this only strips XML-invalid control characters. All Unicode text is preserved:
    em dashes, smart quotes, accented names (José, Olú, García-López), etc.
    """
    if not text:
        return ""
    return _CONTROL_CHAR_RE.sub('', str(text))


# ============================================================
# PUBLIC API
# ============================================================

def generate_cv_docx(cv_data, template="classic"):
    """
    Generate an ATS-friendly resume DOCX.

    Args:
        cv_data: dict with personal, summary, experience, education, skills, certifications
        template: "classic", "modern", or "minimal"

    Returns:
        bytes: DOCX file content
    """
    if template not in TEMPLATES:
        template = "classic"

    personal = cv_data.get("personal", {})
    summary = cv_data.get("summary", "")
    experience = cv_data.get("experience", [])
    education = cv_data.get("education", [])
    raw_skills = cv_data.get("skills", [])
    certifications = cv_data.get("certifications", [])

    skills = _flatten_skills(raw_skills)

    if template == "modern":
        doc = _render_modern(personal, summary, experience, education, skills, certifications)
    elif template == "minimal":
        doc = _render_minimal(personal, summary, experience, education, skills, certifications)
    else:
        doc = _render_classic(personal, summary, experience, education, skills, certifications)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


# ============================================================
# SHARED HELPERS
# ============================================================

def _setup_document(font_name='Calibri', font_size=10, margin_cm=2.54):
    """Create a new Document with standard settings."""
    doc = Document()

    # Set margins
    for section in doc.sections:
        section.top_margin = Cm(margin_cm)
        section.bottom_margin = Cm(margin_cm)
        section.left_margin = Cm(margin_cm)
        section.right_margin = Cm(margin_cm)

    # Set default font
    style = doc.styles['Normal']
    style.font.name = font_name
    style.font.size = Pt(font_size)
    style.font.color.rgb = COLOR_BODY
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.line_spacing = 1.15

    return doc


def _add_paragraph_border_bottom(paragraph, color="1f2937", width="4"):
    """Add a bottom border to a paragraph (used for classic section headers)."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), width)  # 1/8 point units (4 = 0.5pt)
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def _add_tab_stop(paragraph, position_cm, alignment=WD_TAB_ALIGNMENT.RIGHT):
    """Add a tab stop to a paragraph for right-aligned dates."""
    tab_stops = paragraph.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Cm(position_cm), alignment)


def _set_run_font(run, name='Calibri', size=None, bold=False, color=None):
    """Configure a run's font properties."""
    run.font.name = name
    if size:
        run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color


# ============================================================
# TEMPLATE 1: CLASSIC
# Traditional single-column, underline separators, standard corporate.
# ============================================================

def _render_classic(personal, summary, experience, education, skills, certifications):
    doc = _setup_document()

    # -- Name (centered, 18pt bold) --
    name = _docx_safe(personal.get("full_name", ""))
    if name:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(name)
        _set_run_font(run, size=18, bold=True, color=COLOR_DARK)

    # -- Contact line (centered, pipe-separated) --
    contact = _docx_safe(_format_contact_line(personal))
    if contact:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(8)
        run = p.add_run(contact)
        _set_run_font(run, size=10, color=COLOR_GRAY)

    # -- Professional Summary --
    if summary:
        _classic_section_header(doc, "PROFESSIONAL SUMMARY")
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(_docx_safe(summary))
        _set_run_font(run, size=10, color=COLOR_BODY)

    # -- Skills --
    if skills:
        _classic_section_header(doc, "SKILLS")
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(_docx_safe(", ".join(skills)))
        _set_run_font(run, size=10, color=COLOR_BODY)

    # -- Professional Experience --
    if experience:
        _classic_section_header(doc, "PROFESSIONAL EXPERIENCE")
        for i, exp in enumerate(experience):
            title = _docx_safe(exp.get("title", ""))
            company = _docx_safe(exp.get("company", ""))
            dates = _docx_safe(_format_date_range(exp))

            # Title + Company on left, dates right-aligned via tab stop
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(2) if i > 0 else Pt(0)
            p.paragraph_format.space_after = Pt(1)

            title_company = title
            if company:
                title_company += f", {company}"

            run = p.add_run(title_company)
            _set_run_font(run, size=11, bold=True, color=COLOR_DARK)

            if dates:
                _add_tab_stop(p, TEXT_WIDTH_CM)
                run = p.add_run("\t" + dates)
                _set_run_font(run, size=10, color=COLOR_GRAY)

            # Bullets
            for bullet in exp.get("bullets", []):
                bp = doc.add_paragraph(style='List Bullet')
                bp.paragraph_format.space_before = Pt(1)
                bp.paragraph_format.space_after = Pt(1)
                bp.paragraph_format.left_indent = Cm(1.27)
                run = bp.add_run(_docx_safe(bullet))
                _set_run_font(run, size=10, color=COLOR_BODY)

            if i < len(experience) - 1:
                # Small gap between experiences
                sp = doc.add_paragraph()
                sp.paragraph_format.space_before = Pt(4)
                sp.paragraph_format.space_after = Pt(0)

    # -- Education --
    if education:
        _classic_section_header(doc, "EDUCATION")
        for i, edu in enumerate(education):
            degree = _docx_safe(edu.get("degree", ""))
            institution = _docx_safe(edu.get("institution", ""))
            grad_date = _docx_safe(edu.get("graduation_date", ""))
            details = _docx_safe(edu.get("details", ""))

            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(2) if i > 0 else Pt(0)
            p.paragraph_format.space_after = Pt(1)

            deg_inst = degree
            if institution:
                deg_inst += f", {institution}"

            run = p.add_run(deg_inst)
            _set_run_font(run, size=11, bold=True, color=COLOR_DARK)

            if grad_date:
                _add_tab_stop(p, TEXT_WIDTH_CM)
                run = p.add_run("\t" + grad_date)
                _set_run_font(run, size=10, color=COLOR_GRAY)

            if details:
                dp = doc.add_paragraph()
                dp.paragraph_format.left_indent = Cm(0.5)
                dp.paragraph_format.space_after = Pt(2)
                run = dp.add_run(details)
                _set_run_font(run, size=10, color=COLOR_BODY)

    # -- Certifications --
    if certifications:
        _classic_section_header(doc, "CERTIFICATIONS")
        for cert in certifications:
            cert_name = _docx_safe(cert.get("name", ""))
            issuer = _docx_safe(cert.get("issuer", ""))
            date = _docx_safe(cert.get("date", ""))
            line = cert_name
            if issuer:
                line += f" — {issuer}"
            if date:
                line += f" ({date})"
            bp = doc.add_paragraph(style='List Bullet')
            bp.paragraph_format.space_before = Pt(1)
            bp.paragraph_format.space_after = Pt(1)
            bp.paragraph_format.left_indent = Cm(1.27)
            run = bp.add_run(line)
            _set_run_font(run, size=10, color=COLOR_BODY)

    return doc


def _classic_section_header(doc, title):
    """Section header with full-width bottom border (underline separator)."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(_docx_safe(title))
    _set_run_font(run, size=13, bold=True, color=COLOR_DARK)
    _add_paragraph_border_bottom(p, color="1f2937", width="4")


# ============================================================
# TEMPLATE 2: MODERN
# Blue accent line, uppercase headers, compact but readable.
# Uses single-column tables for the blue left border effect.
# ============================================================

def _render_modern(personal, summary, experience, education, skills, certifications):
    doc = _setup_document()

    # -- Name (bold, blue, left-aligned) --
    name = _docx_safe(personal.get("full_name", ""))
    if name:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(name)
        _set_run_font(run, size=20, bold=True, color=COLOR_BLUE)

    # -- Contact (stacked on 2 lines) --
    contact_parts = []
    if personal.get("email"):
        contact_parts.append(personal["email"])
    if personal.get("phone"):
        contact_parts.append(personal["phone"])
    if personal.get("location"):
        contact_parts.append(personal["location"])

    line2_parts = []
    if personal.get("linkedin"):
        line2_parts.append(personal["linkedin"])
    if personal.get("portfolio"):
        line2_parts.append(personal["portfolio"])

    if contact_parts:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(1)
        run = p.add_run(_docx_safe(" | ".join(contact_parts)))
        _set_run_font(run, size=10, color=COLOR_GRAY)

    if line2_parts:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(8)
        run = p.add_run(_docx_safe(" | ".join(line2_parts)))
        _set_run_font(run, size=10, color=COLOR_GRAY)
    elif contact_parts:
        # Add spacing after contact if no second line
        doc.paragraphs[-1].paragraph_format.space_after = Pt(8)

    # -- Professional Summary --
    if summary:
        _modern_section_header(doc, "PROFESSIONAL SUMMARY")
        _modern_accent_block(doc, [(_docx_safe(summary), 10, False, COLOR_BODY)])

    # -- Skills --
    if skills:
        _modern_section_header(doc, "SKILLS")
        _modern_accent_block(doc, [(_docx_safe(", ".join(skills)), 10, False, COLOR_BODY)])

    # -- Experience --
    if experience:
        _modern_section_header(doc, "EXPERIENCE")
        for i, exp in enumerate(experience):
            title = _docx_safe(exp.get("title", ""))
            company = _docx_safe(exp.get("company", ""))
            dates = _docx_safe(_format_date_range(exp))

            header_line = title
            if company:
                header_line += f" | {company}"

            lines = [(header_line, 10, True, COLOR_DARK)]
            if dates:
                lines.append((dates, 10, False, COLOR_GRAY))

            for bullet in exp.get("bullets", []):
                lines.append(("• " + _docx_safe(bullet), 10, False, COLOR_BODY))

            _modern_accent_block(doc, lines)

            if i < len(experience) - 1:
                sp = doc.add_paragraph()
                sp.paragraph_format.space_before = Pt(2)
                sp.paragraph_format.space_after = Pt(0)

    # -- Education --
    if education:
        _modern_section_header(doc, "EDUCATION")
        for edu in education:
            degree = _docx_safe(edu.get("degree", ""))
            institution = _docx_safe(edu.get("institution", ""))
            grad_date = _docx_safe(edu.get("graduation_date", ""))
            details = _docx_safe(edu.get("details", ""))

            deg_line = degree
            if institution:
                deg_line += f" | {institution}"

            lines = [(deg_line, 10, True, COLOR_DARK)]
            if grad_date:
                lines.append((grad_date, 10, False, COLOR_GRAY))
            if details:
                lines.append((details, 10, False, COLOR_BODY))

            _modern_accent_block(doc, lines)

    # -- Certifications --
    if certifications:
        _modern_section_header(doc, "CERTIFICATIONS")
        for cert in certifications:
            cert_name = _docx_safe(cert.get("name", ""))
            issuer = _docx_safe(cert.get("issuer", ""))
            date = _docx_safe(cert.get("date", ""))
            line = cert_name
            if issuer:
                line += f" — {issuer}"
            if date:
                line += f" ({date})"
            _modern_accent_block(doc, [("• " + line, 10, False, COLOR_BODY)])

    return doc


def _modern_section_header(doc, title):
    """Uppercase section header in blue."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(_docx_safe(title))
    _set_run_font(run, size=13, bold=True, color=COLOR_BLUE)


def _modern_accent_block(doc, lines):
    """Render content inside a single-column table with a blue left border.

    Args:
        doc: Document instance
        lines: list of (text, font_size, bold, color) tuples
    """
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Remove all borders, then add only left border in blue
    cell = table.cell(0, 0)
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')

    for border_name in ['top', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'none')
        border.set(qn('w:sz'), '0')
        border.set(qn('w:space'), '0')
        tcBorders.append(border)

    left_border = OxmlElement('w:left')
    left_border.set(qn('w:val'), 'single')
    left_border.set(qn('w:sz'), '4')  # 0.5pt
    left_border.set(qn('w:space'), '0')
    left_border.set(qn('w:color'), '2563eb')  # Blue
    tcBorders.append(left_border)
    tcPr.append(tcBorders)

    # Set cell padding
    tcMar = OxmlElement('w:tcMar')
    for side, val in [('top', '40'), ('bottom', '40'), ('left', '120'), ('right', '0')]:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:w'), val)
        el.set(qn('w:type'), 'dxa')
        tcMar.append(el)
    tcPr.append(tcMar)

    # Remove the default empty paragraph
    for paragraph in cell.paragraphs:
        p_element = paragraph._element
        p_element.getparent().remove(p_element)

    # Add content lines
    for text, font_size, bold, color in lines:
        p = cell.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        run = p.add_run(text)
        _set_run_font(run, size=font_size, bold=bold, color=color)


# ============================================================
# TEMPLATE 3: MINIMAL
# Ultra-clean, generous whitespace, no separators or bullets.
# ============================================================

def _render_minimal(personal, summary, experience, education, skills, certifications):
    doc = _setup_document(margin_cm=2.54)

    # Wider margins for minimal feel
    for section in doc.sections:
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(3.0)

    # Recalculate text width for wider margins: 21 - 3.0 - 3.0 = 15.0cm
    minimal_text_width = 15.0

    # -- Name (regular weight, centered) --
    name = _docx_safe(personal.get("full_name", ""))
    if name:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(name)
        _set_run_font(run, size=18, bold=False, color=COLOR_DARK)

    # -- Contact (centered, pipe-separated) --
    contact = _docx_safe(_format_contact_line(personal, " | "))
    if contact:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(16)
        run = p.add_run(contact)
        _set_run_font(run, size=9, color=COLOR_GRAY)

    # -- Summary --
    if summary:
        _minimal_section_header(doc, "Summary")
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(10)
        run = p.add_run(_docx_safe(summary))
        _set_run_font(run, size=10, color=COLOR_BODY)

    # -- Skills --
    if skills:
        _minimal_section_header(doc, "Skills")
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(10)
        run = p.add_run(_docx_safe(", ".join(skills)))
        _set_run_font(run, size=10, color=COLOR_BODY)

    # -- Experience --
    if experience:
        _minimal_section_header(doc, "Experience")
        for i, exp in enumerate(experience):
            title = _docx_safe(exp.get("title", ""))
            company = _docx_safe(exp.get("company", ""))
            dates = _docx_safe(_format_date_range(exp))

            # Title bold
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(1)
            run = p.add_run(title)
            _set_run_font(run, size=10, bold=True, color=COLOR_DARK)

            # Company | Dates
            sub_parts = []
            if company:
                sub_parts.append(company)
            if dates:
                sub_parts.append(dates)
            if sub_parts:
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(3)
                run = p.add_run(" | ".join(sub_parts))
                _set_run_font(run, size=10, color=COLOR_GRAY)

            # Indented text blocks (no bullet symbols — minimal design)
            for bullet in exp.get("bullets", []):
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(0.5)
                p.paragraph_format.space_before = Pt(1)
                p.paragraph_format.space_after = Pt(2)
                run = p.add_run(_docx_safe(bullet))
                _set_run_font(run, size=10, color=COLOR_BODY)

            if i < len(experience) - 1:
                sp = doc.add_paragraph()
                sp.paragraph_format.space_before = Pt(6)
                sp.paragraph_format.space_after = Pt(0)

    # -- Education --
    if education:
        _minimal_section_header(doc, "Education")
        for edu in education:
            degree = _docx_safe(edu.get("degree", ""))
            institution = _docx_safe(edu.get("institution", ""))
            grad_date = _docx_safe(edu.get("graduation_date", ""))
            details = _docx_safe(edu.get("details", ""))

            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(1)
            run = p.add_run(degree)
            _set_run_font(run, size=10, bold=True, color=COLOR_DARK)

            sub_parts = []
            if institution:
                sub_parts.append(institution)
            if grad_date:
                sub_parts.append(grad_date)
            if sub_parts:
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(2)
                run = p.add_run(" | ".join(sub_parts))
                _set_run_font(run, size=10, color=COLOR_GRAY)

            if details:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(0.5)
                p.paragraph_format.space_after = Pt(4)
                run = p.add_run(details)
                _set_run_font(run, size=10, color=COLOR_BODY)

    # -- Certifications --
    if certifications:
        _minimal_section_header(doc, "Certifications")
        for cert in certifications:
            cert_name = _docx_safe(cert.get("name", ""))
            issuer = _docx_safe(cert.get("issuer", ""))
            date = _docx_safe(cert.get("date", ""))
            line = cert_name
            if issuer:
                line += f" — {issuer}"
            if date:
                line += f" ({date})"
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(line)
            _set_run_font(run, size=10, color=COLOR_BODY)

    return doc


def _minimal_section_header(doc, title):
    """Simple bold header with generous spacing, no separator."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(_docx_safe(title))
    _set_run_font(run, size=12, bold=True, color=COLOR_DARK)
