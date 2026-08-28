"""
Resume Parser Module
Handles extraction of text from PDF, DOCX, and plain text resumes.
"""

import os
import re
import zipfile
from pypdf import PdfReader
from docx import Document


# ---------------------------------------------------------------------------
# Upload safety limits (AI-02)
#
# The 5MB upload cap bounds what arrives, not what it expands into. A small
# PDF can carry tens of thousands of pages, and a DOCX is a zip archive that
# can inflate to gigabytes. Both turn a single upload into a CPU/memory DoS,
# so the parsers enforce their own ceilings.
# ---------------------------------------------------------------------------
MAX_PDF_PAGES = 50
# 10MB, not 50. The guard stops bombs either way, but everything under the
# ceiling still gets handed to lxml, whose DOM costs far more than the bytes:
# a legitimate 2.4MB upload declaring 31MB measured ~440MB and 10s in
# python-docx. 10MB keeps >12x headroom over the largest DOCX this product
# generates (~0.8MB expanded) while cutting that worst case by roughly a third.
MAX_DOCX_UNCOMPRESSED = 10 * 1024 * 1024   # 10MB expanded
MAX_DOCX_COMPRESSION_RATIO = 150           # expanded / compressed
MAX_DOCX_ENTRIES = 2000                    # a real DOCX has tens, not thousands

# Only the compression methods Word actually produces. See _docx_zip_is_safe.
_ALLOWED_COMPRESSION = (zipfile.ZIP_STORED, zipfile.ZIP_DEFLATED)

# The PDF spec allows the header anywhere in the first 1024 bytes, and real
# files exploit that — a UTF-8 BOM or a stray newline before "%PDF-" is common
# and every reader accepts it. Scanning the window instead of byte 0 keeps us
# from rejecting valid resumes.
_PDF_HEADER_WINDOW = 1024

# Leading bytes that identify each accepted format. The upload routes check the
# filename extension, which the client controls; this checks the actual content.
_MAGIC = {
    "pdf": b"%PDF-",
    "docx": b"PK\x03\x04",   # DOCX is a zip archive
}


def sniff_matches(file_path, file_type):
    """
    Return True if the file's leading bytes match the claimed type.

    A renamed file (evil.zip -> resume.pdf) passes an extension check but
    fails here, so the parser never hands unexpected content to a library.
    """
    magic = _MAGIC.get(file_type)
    if not magic:
        return False
    try:
        with open(file_path, "rb") as fh:
            head = fh.read(_PDF_HEADER_WINDOW)
    except Exception:
        return False

    if file_type == "pdf":
        # Header may sit anywhere in the first 1024 bytes (see above).
        return magic in head
    # DOCX is a zip: the local file header must be at byte 0.
    return head.startswith(magic)


def extract_text_from_pdf(file_path):
    """Extract text from a PDF file."""
    if not sniff_matches(file_path, "pdf"):
        return None, "That file doesn't look like a valid PDF. Please upload a real PDF or paste your resume text."

    try:
        reader = PdfReader(file_path)

        page_count = len(reader.pages)
        if page_count > MAX_PDF_PAGES:
            return None, f"That PDF has {page_count} pages (limit {MAX_PDF_PAGES}). Please upload just your resume."

        text_parts = []
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)

        full_text = "\n".join(text_parts)

        if not full_text.strip():
            return None, "The PDF appears to be empty or contains only images/scanned content. Try pasting your resume text directly."

        return clean_text(full_text), None

    except Exception as e:
        # Log the library detail, but don't hand it to the client.
        print(f"PDF parse error: {type(e).__name__}: {e}")
        if "decrypt" in str(e).lower() or "password" in str(e).lower():
            return None, "That PDF is password-protected. Please remove the password or paste your resume text."
        return None, "Could not read that PDF. Please try a different file or paste your resume text."


def _docx_zip_is_safe(file_path):
    """
    Check a DOCX archive before python-docx opens it.

    Two stages, because the first one is cheap but not trustworthy:

    1. The central directory's declared sizes. Free to read, but every byte of
       it is attacker-controlled, so a crafted archive can declare a harmless
       100-byte member while carrying megabytes. Used only as a fast reject.
    2. An actual bounded read of each member. This is the load-bearing check:
       it counts real decompressed bytes and stops the moment the budget is
       exceeded, so it holds regardless of what the metadata claims.

    Returns (True, None) if safe, else (False, user-facing message).
    """
    try:
        with zipfile.ZipFile(file_path) as zf:
            infos = zf.infolist()

            if len(infos) > MAX_DOCX_ENTRIES:
                return False, "That DOCX has too many parts to process safely. Please paste your resume text instead."

            # Stage 1 — cheap reject on declared sizes.
            declared = sum(i.file_size for i in infos)
            compressed = sum(i.compress_size for i in infos) or 1
            if declared > MAX_DOCX_UNCOMPRESSED:
                return False, "That DOCX expands to more than we can process. Please upload a smaller file or paste your resume text."
            if declared / compressed > MAX_DOCX_COMPRESSION_RATIO:
                return False, "That DOCX looks malformed and can't be processed safely. Please paste your resume text instead."

            # Word, LibreOffice and python-docx only ever emit STORED or
            # DEFLATE. Reject anything else *before* reading a byte: CPython
            # decompresses BZIP2/LZMA members with no length bound
            # (ZipExtFile._read1 calls decompress(data) without max_length), so
            # a chunked read of one cannot be kept small. Measured: a 36KB
            # DOCX with a BZIP2 member allocated 404MB inside this function.
            for info in infos:
                if info.compress_type not in _ALLOWED_COMPRESSION:
                    return False, "That DOCX uses an unsupported compression method. Please re-save it in Word or paste your resume text."

            # Stage 2 — count real decompressed bytes, capped.
            #
            # zipfile truncates each member's output at its DECLARED file_size
            # (ZipExtFile._read1 does `data = data[:self._left]`), so a reader
            # that trusts the declaration can never see more than the attacker
            # claims — the loop would count to the fake 100 and pass. Raising
            # the declared size first makes the real stream observable. Memory
            # stays bounded because DEFLATE decompression is chunk-limited by
            # max_length, and we abort the moment the budget is exceeded.
            actual = 0
            for info in infos:
                info.file_size = MAX_DOCX_UNCOMPRESSED + 1
                with zf.open(info) as member:
                    while True:
                        chunk = member.read(64 * 1024)
                        if not chunk:
                            break
                        actual += len(chunk)
                        if actual > MAX_DOCX_UNCOMPRESSED:
                            return False, "That DOCX expands to more than we can process. Please upload a smaller file or paste your resume text."

        return True, None
    except zipfile.BadZipFile:
        return False, "That file doesn't look like a valid DOCX. Please upload a real DOCX or paste your resume text."
    except Exception:
        return False, "Could not read that DOCX file. Please try a different file or paste your resume text."


def extract_text_from_docx(file_path):
    """Extract text from a DOCX file."""
    if not sniff_matches(file_path, "docx"):
        return None, "That file doesn't look like a valid DOCX. Please upload a real DOCX or paste your resume text."

    safe, sniff_error = _docx_zip_is_safe(file_path)
    if not safe:
        return None, sniff_error

    try:
        doc = Document(file_path)
        text_parts = []

        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)

        # Also extract text from tables (common in resumes)
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_parts.append(" | ".join(row_text))

        full_text = "\n".join(text_parts)

        if not full_text.strip():
            return None, "The DOCX file appears to be empty. Try pasting your resume text directly."

        return clean_text(full_text), None

    except Exception as e:
        print(f"DOCX parse error: {type(e).__name__}: {e}")
        return None, "Could not read that DOCX. Please try a different file or paste your resume text."


def extract_text_from_paste(text):
    """Process pasted resume text."""
    if not text or not text.strip():
        return None, "No resume text provided. Please paste your resume content."

    cleaned = clean_text(text)

    if len(cleaned.split()) < 20:
        return None, "The pasted text seems too short to be a resume. Please paste your full resume content."

    return cleaned, None


def clean_text(text):
    """Clean and normalize extracted text."""
    # Replace multiple whitespace with single space (but preserve newlines)
    text = re.sub(r'[^\S\n]+', ' ', text)
    # Replace multiple consecutive newlines with double newline
    text = re.sub(r'\n{3,}', '\n\n', text)
    # Remove leading/trailing whitespace from each line
    lines = [line.strip() for line in text.split('\n')]
    text = '\n'.join(lines)
    # Remove leading/trailing whitespace from the whole text
    text = text.strip()
    return text


def parse_resume(file_path=None, pasted_text=None, file_type=None):
    """
    Main entry point for resume parsing.

    Args:
        file_path: Path to uploaded file (PDF or DOCX)
        pasted_text: Directly pasted resume text
        file_type: Type of file ('pdf', 'docx', or 'text')

    Returns:
        dict with 'text' (extracted text) and 'error' (error message if any)
    """
    if pasted_text:
        text, error = extract_text_from_paste(pasted_text)
    elif file_path and file_type:
        if file_type == 'pdf':
            text, error = extract_text_from_pdf(file_path)
        elif file_type == 'docx':
            text, error = extract_text_from_docx(file_path)
        else:
            text, error = None, f"Unsupported file type: {file_type}"
    else:
        text, error = None, "No resume provided. Please upload a file or paste your resume text."

    if error:
        return {"text": None, "error": error, "word_count": 0}

    word_count = len(text.split())

    return {
        "text": text,
        "error": None,
        "word_count": word_count
    }
